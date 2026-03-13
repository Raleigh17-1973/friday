"""Word document generator using python-docx."""
from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt

from packages.docgen.generators.base import DocumentContent, DocumentGenerator, DocumentSection


class DocxGenerator(DocumentGenerator):
    """Generate .docx files from DocumentContent."""

    @property
    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def file_extension(self) -> str:
        return "docx"

    def generate(
        self, content: DocumentContent, template_path: str | None = None
    ) -> bytes:
        doc = Document(template_path) if template_path else Document()

        # Title
        doc.add_heading(content.title, level=0)

        # Metadata line
        author = content.metadata.get("author", "")
        date = content.metadata.get("date", "")
        if author or date:
            meta_parts = []
            if author:
                meta_parts.append(f"Author: {author}")
            if date:
                meta_parts.append(f"Date: {date}")
            doc.add_paragraph(" | ".join(meta_parts)).style = "Subtitle"

        for section in content.sections:
            self._add_section(doc, section)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    def _add_section(self, doc: Document, section: DocumentSection) -> None:
        heading_level = min(max(section.level, 1), 9)
        doc.add_heading(section.heading, level=heading_level)

        # Parse body with simple markdown support
        self._add_markdown_body(doc, section.body)

        # Table
        if section.table:
            self._add_table(doc, section.table)

    def _add_markdown_body(self, doc: Document, body: str) -> None:
        if not body:
            return

        for line in body.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                self._add_formatted_runs(p, stripped[2:])
            else:
                p = doc.add_paragraph()
                self._add_formatted_runs(p, stripped)

    def _add_formatted_runs(self, paragraph, text: str) -> None:
        """Handle **bold** and *italic* inline formatting."""
        # Regex to split on bold/italic markers
        pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")
        parts = pattern.split(text)

        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    def _add_table(self, doc: Document, table_data: list[list[str]]) -> None:
        if not table_data:
            return

        rows = len(table_data)
        cols = max(len(r) for r in table_data)
        tbl = doc.add_table(rows=rows, cols=cols, style="Table Grid")

        for r_idx, row in enumerate(table_data):
            for c_idx, cell_text in enumerate(row):
                tbl.rows[r_idx].cells[c_idx].text = cell_text

        # Bold the header row
        if rows > 0:
            for cell in tbl.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
