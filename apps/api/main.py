from __future__ import annotations

import os
import pathlib
from dataclasses import asdict
from typing import Any, Optional

# Load .env from project root before anything else reads os.getenv()
_ROOT = pathlib.Path(__file__).resolve().parents[2]
_dotenv = _ROOT / ".env"
if _dotenv.exists():
    with _dotenv.open() as _f:
        for _line in _f:
            _line = _line.strip()
            if _line and not _line.startswith("#") and "=" in _line:
                _k, _, _v = _line.partition("=")
                os.environ.setdefault(_k.strip(), _v.strip())

from apps.api.service import FridayService
from apps.api.security import AdminAuth, RateLimiter

try:
    from fastapi import FastAPI, File, HTTPException, Request, UploadFile
    from fastapi.middleware.cors import CORSMiddleware
    from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
    from fastapi.staticfiles import StaticFiles
    from pydantic import BaseModel
except ModuleNotFoundError as exc:  # pragma: no cover
    raise RuntimeError(
        "FastAPI is required for apps/api/main.py. Install project dependencies first."
    ) from exc

# In-process upload context store: context_id -> {"filename": str, "text": str, "type": str}
# Lost on restart; use a persistent store (Redis, DB) for production.
_UPLOAD_STORE: dict[str, dict] = {}

app = FastAPI(title="Friday API", version="0.2.0")

# CORS — allow Next.js dev server and same-origin requests
_raw_origins = os.getenv("FRIDAY_ALLOWED_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000,http://127.0.0.1:8000")
_allowed_origins = [o.strip() for o in _raw_origins.split(",") if o.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_UI_DIR = pathlib.Path(__file__).parent.parent.parent / "ui"
if _UI_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(_UI_DIR / "assets")), name="static")

    @app.get("/")
    def serve_ui():
        return FileResponse(str(_UI_DIR / "index.html"))

service = FridayService()
admin_auth = AdminAuth()
rate_limiter = RateLimiter()


# --- Pydantic request models ---

class ChatPayload(BaseModel):
    message: str
    user_id: str = "user-1"
    org_id: str = "org-1"
    conversation_id: str = "conv-1"
    context_packet: dict[str, Any] = {}
    # Optional context IDs from POST /upload — injected as context blocks
    context_ids: list[str] = []
    # Optional workspace context — injected into planning message when set
    workspace_id: Optional[str] = None


class FeedbackPayload(BaseModel):
    approved: bool
    notes: str = ""


class PromoteCandidatePayload(BaseModel):
    candidate_id: str
    approved: bool = False


class ScaffoldAgentPayload(BaseModel):
    id: str
    name: str
    purpose: str
    owner: str


class EvalPayload(BaseModel):
    suite: str = "core-routing"


class ProcessCreatePayload(BaseModel):
    org_id: str = "org-1"
    process_name: str
    trigger: str = ""
    steps: list[dict] = []
    decision_points: list[dict] = []
    roles: list[str] = []
    tools: list[str] = []
    exceptions: list[dict] = []
    kpis: list[dict] = []
    mermaid_flowchart: str = ""
    mermaid_swimlane: str = ""
    status: str = "draft"


class ProcessUpdatePayload(BaseModel):
    bump: str = "patch"           # major | minor | patch
    author: str = "user"
    changelog_entry: str = ""
    # fields to update (any subset)
    process_name: Optional[str] = None
    trigger: Optional[str] = None
    steps: Optional[list] = None
    decision_points: Optional[list] = None
    roles: Optional[list] = None
    tools: Optional[list] = None
    exceptions: Optional[list] = None
    kpis: Optional[list] = None
    mermaid_flowchart: Optional[str] = None
    mermaid_swimlane: Optional[str] = None
    status: Optional[str] = None


@app.middleware("http")
async def add_security_and_limits(request: Request, call_next):
    if len(request.url.path) > 2048:
        return JSONResponse({"detail": "path too long"}, status_code=414)

    key = request.client.host if request.client else "unknown"
    rate_limiter.check(key)
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "no-referrer"
    response.headers["Cache-Control"] = "no-store"
    return response


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.get("/ready")
def ready() -> dict:
    """Readiness probe — checks that all critical dependencies are reachable.

    Used by Fly.io health checks and Kubernetes readiness probes.
    Returns 200 when the service can accept traffic, 503 when not ready.
    """
    checks: dict[str, str] = {}
    ok = True

    # Memory / DB connectivity
    try:
        # Lightweight check: just verify the memory service is initialised
        _ = service.memory  # would raise if initialisation failed
        checks["memory"] = "ok"
    except Exception as exc:
        checks["memory"] = f"error: {exc}"
        ok = False

    # Agent registry
    try:
        agents = service.registry.list_active()
        checks["registry"] = f"ok ({len(agents)} agents)"
    except Exception as exc:
        checks["registry"] = f"error: {exc}"
        ok = False

    # LLM provider (advisory — does not block readiness)
    checks["llm"] = "connected" if service.llm is not None else "offline (stub mode)"

    status_code = 200 if ok else 503
    return JSONResponse(
        content={"status": "ready" if ok else "degraded", "checks": checks},
        status_code=status_code,
    )


@app.post("/chat")
def chat(payload: ChatPayload) -> dict:
    try:
        result = service.execute_chat_payload(payload.model_dump(), upload_store=_UPLOAD_STORE)
        # Persist conversation messages
        try:
            msg = payload.message.strip()
            if msg:
                service.conversations.add_message(
                    thread_id=payload.conversation_id,
                    role="user",
                    content=msg,
                    metadata={"org_id": payload.org_id, "workspace_id": payload.workspace_id},
                )
                friday_text = str((result.get("final_answer") or {}).get("direct_answer") or result.get("response") or "")
                if friday_text:
                    meta: dict = {"org_id": payload.org_id}
                    if result.get("run_id"): meta["run_id"] = result["run_id"]
                    if result.get("write_actions"): meta["write_actions"] = result["write_actions"]
                    service.conversations.add_message(
                        thread_id=payload.conversation_id,
                        role="friday",
                        content=friday_text,
                        metadata=meta,
                    )
        except Exception:
            pass  # best-effort
        return result
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.post("/chat/stream")
def chat_stream(payload: ChatPayload):
    """Server-sent events endpoint that streams synthesis tokens as they arrive from the LLM.

    SSE event types:
      event: status  — pipeline stage label
      event: token   — one synthesis token chunk
      event: done    — final metadata JSON
      event: error   — error message
    """
    import json as _json
    from packages.common.models import ChatRequest

    message = payload.message.strip()
    if not message:
        raise HTTPException(status_code=400, detail="message is required")

    request = ChatRequest(
        user_id=payload.user_id,
        org_id=payload.org_id,
        conversation_id=payload.conversation_id,
        message=message,
        context_packet=payload.context_packet,
        workspace_id=payload.workspace_id,
    )

    def _event_generator():
        full_text_parts: list[str] = []
        done_evt: dict | None = None
        for evt in service.manager.run_streaming(request):
            event_type = evt.get("event", "message")
            if event_type == "token":
                full_text_parts.append(evt.get("text", ""))
            if event_type == "done":
                done_evt = evt
                # Attempt doc generation if this was a full_deliverable request
                if evt.get("output_format") == "full_deliverable" and service.docgen is not None:
                    full_text = "".join(full_text_parts)
                    response_stub = {
                        "planner": {"output_format": "full_deliverable", "problem_statement": message},
                        "final_answer": {"direct_answer": full_text, "artifacts": {}},
                    }
                    service._maybe_generate_document(message, payload.org_id, response_stub)
                    if response_stub.get("generated_document"):
                        evt = {**evt, "generated_document": response_stub["generated_document"]}
                        done_evt = evt
            yield f"event: {event_type}\ndata: {_json.dumps(evt)}\n\n"
        # Persist conversation messages after streaming completes
        if full_text_parts:
            try:
                service.conversations.add_message(
                    thread_id=payload.conversation_id,
                    role="user",
                    content=message,
                    metadata={"org_id": payload.org_id, "workspace_id": payload.workspace_id},
                )
                friday_text = "".join(full_text_parts)
                meta: dict = {"org_id": payload.org_id}
                if done_evt:
                    if done_evt.get("run_id"): meta["run_id"] = done_evt["run_id"]
                    if done_evt.get("write_actions"): meta["write_actions"] = done_evt["write_actions"]
                service.conversations.add_message(
                    thread_id=payload.conversation_id,
                    role="friday",
                    content=friday_text,
                    metadata=meta,
                )
            except Exception:
                pass  # persistence is best-effort — never break the stream

    return StreamingResponse(
        _event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache, no-transform", "X-Accel-Buffering": "no"},
    )


# ── Conversation endpoints ─────────────────────────────────────────────────────

class ConversationCreate(BaseModel):
    org_id: str = "org-1"
    workspace_id: Optional[str] = None
    title: str = "New conversation"
    thread_id: Optional[str] = None


class ConversationRename(BaseModel):
    title: str


@app.get("/conversations")
def list_conversations(org_id: str = "org-1") -> list[dict]:
    return [t.to_dict() for t in service.conversations.list_threads(org_id=org_id)]


@app.post("/conversations", status_code=201)
def create_conversation(payload: ConversationCreate) -> dict:
    thread = service.conversations.create_thread(
        org_id=payload.org_id,
        workspace_id=payload.workspace_id,
        title=payload.title,
        thread_id=payload.thread_id,
    )
    return thread.to_dict()


@app.get("/conversations/{thread_id}/messages")
def get_conversation_messages(thread_id: str) -> list[dict]:
    messages = service.conversations.get_messages(thread_id)
    return [m.to_dict() for m in messages]


@app.patch("/conversations/{thread_id}")
def rename_conversation(thread_id: str, payload: ConversationRename) -> dict:
    thread = service.conversations.rename_thread(thread_id, payload.title)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    return thread.to_dict()


@app.delete("/conversations/{thread_id}", status_code=204)
def delete_conversation(thread_id: str) -> None:
    service.conversations.delete_thread(thread_id)
    try:
        service.memory.clear_conversation(thread_id)
    except Exception:
        pass


# ──────────────────────────────────────────────────────────────────────────────

@app.post("/runs")
def create_run(payload: ChatPayload) -> dict:
    def _execute(run_payload: dict) -> dict:
        return chat(ChatPayload(**run_payload))

    record = service.workflow.run(payload.model_dump(), _execute)
    return {
        "workflow_id": record.workflow_id,
        "status": record.status,
        "result": record.result,
    }


@app.get("/runs/{run_id}")
def get_run(run_id: str) -> dict:
    workflow = service.workflow.get(run_id)
    if workflow is not None:
        return {
            "workflow_id": workflow.workflow_id,
            "status": workflow.status,
            "result": workflow.result,
        }

    trace = service.audit.get_run(run_id)
    if trace is None:
        raise HTTPException(status_code=404, detail="run not found")
    return asdict(trace)


@app.post("/runs/{run_id}/feedback")
def run_feedback(run_id: str, payload: FeedbackPayload) -> dict:
    """Record user approval or rejection of a run output.

    The episodic learning loop uses this signal to bias future planning toward
    specialist combinations that have been approved for similar requests.
    """
    try:
        service.manager.record_feedback(run_id, approved=payload.approved, notes=payload.notes)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    return {"status": "recorded", "run_id": run_id, "approved": payload.approved}


@app.post("/approvals/{approval_id}/approve")
def approve(approval_id: str) -> dict:
    try:
        req = service.approvals.approve(approval_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="approval not found") from exc
    return asdict(req)


@app.post("/approvals/{approval_id}/reject")
def reject(approval_id: str) -> dict:
    try:
        req = service.approvals.reject(approval_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="approval not found") from exc
    return asdict(req)


@app.get("/memories")
def list_memories(org_id: str = "org-1") -> dict:
    """Return all stored semantic memories for an org as a flat key/value list."""
    semantic: dict = service.memory._semantic_by_org.get(org_id, {})
    memories = [{"key": k, "value": str(v)} for k, v in semantic.items()]
    return {"org_id": org_id, "memories": memories, "count": len(memories)}


@app.delete("/memories/{org_id}/{key}")
def delete_memory(org_id: str, key: str) -> dict:
    """Remove a single semantic memory entry."""
    store = service.memory._semantic_by_org.get(org_id, {})
    if key not in store:
        raise HTTPException(status_code=404, detail="memory key not found")
    del service.memory._semantic_by_org[org_id][key]
    return {"deleted": key}


@app.get("/memories/search")
def search_memories(org_id: str, q: str) -> dict:
    return service.memory.search(org_id=org_id, query=q)


@app.post("/memories/candidates/promote")
def promote_memory_candidate(payload: PromoteCandidatePayload) -> dict:
    try:
        return service.memory.promote_candidate(payload.candidate_id, approved=payload.approved)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="candidate not found") from exc


@app.get("/approvals")
def list_approvals(status: str = "pending") -> dict:
    if status == "all":
        items = service.approvals.list_all()
    else:
        items = service.approvals.list_pending()
    return {"approvals": [asdict(r) for r in items]}


# ── Tasks ──────────────────────────────────────────────────────────────────────

class TaskCreate(BaseModel):
    title: str
    description: str = ""
    assignee: Optional[str] = None
    due_date: Optional[str] = None
    priority: str = "medium"
    status: str = "open"
    workspace_id: Optional[str] = None
    okr_id: Optional[str] = None
    kr_id: Optional[str] = None
    process_id: Optional[str] = None
    initiative_id: Optional[str] = None
    created_by: str = "system"


class TaskUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    assignee: Optional[str] = None
    due_date: Optional[str] = None
    priority: Optional[str] = None
    status: Optional[str] = None
    workspace_id: Optional[str] = None
    okr_id: Optional[str] = None
    kr_id: Optional[str] = None
    process_id: Optional[str] = None
    initiative_id: Optional[str] = None


@app.get("/tasks")
def list_tasks(
    assignee: Optional[str] = None,
    workspace_id: Optional[str] = None,
    status: Optional[str] = None,
    priority: Optional[str] = None,
    due_before: Optional[str] = None,
    okr_id: Optional[str] = None,
    limit: int = 200,
) -> list[dict]:
    tasks = service.tasks.list(
        assignee=assignee,
        workspace_id=workspace_id,
        status=status,
        priority=priority,
        due_before=due_before,
        okr_id=okr_id,
        limit=limit,
    )
    return [t.to_dict() for t in tasks]


@app.post("/tasks", status_code=201)
def create_task(payload: TaskCreate) -> dict:
    task = service.tasks.create(
        title=payload.title,
        description=payload.description,
        assignee=payload.assignee,
        due_date=payload.due_date,
        priority=payload.priority,
        status=payload.status,
        workspace_id=payload.workspace_id,
        okr_id=payload.okr_id,
        kr_id=payload.kr_id,
        process_id=payload.process_id,
        initiative_id=payload.initiative_id,
        created_by=payload.created_by,
    )
    # Notify assignee when task is assigned to someone
    if task.assignee:
        try:
            service.notifications.create(
                recipient_id=task.assignee,
                title=f"Task assigned: {task.title}",
                body=f"Priority: {task.priority}" + (f" · Due {task.due_date}" if task.due_date else ""),
                type="task_assigned",
                entity_type="task",
                entity_id=task.task_id,
            )
        except Exception:
            pass
    # Log to activity feed
    try:
        service.activity.log(
            action="task.created",
            entity_type="task",
            entity_id=task.task_id,
            entity_title=task.title,
            actor_id=payload.created_by or "system",
            assignee=task.assignee or "",
            priority=task.priority,
        )
    except Exception:
        pass
    return task.to_dict()


@app.get("/tasks/overdue")
def list_overdue_tasks() -> list[dict]:
    return [t.to_dict() for t in service.tasks.overdue()]


@app.get("/tasks/due-soon")
def list_due_soon_tasks(days: int = 7) -> list[dict]:
    return [t.to_dict() for t in service.tasks.due_soon(days=days)]


@app.get("/tasks/{task_id}")
def get_task(task_id: str) -> dict:
    task = service.tasks.get(task_id)
    if task is None:
        raise HTTPException(status_code=404, detail="Task not found")
    return task.to_dict()


@app.put("/tasks/{task_id}")
def update_task(task_id: str, payload: TaskUpdate) -> dict:
    changes = {k: v for k, v in payload.model_dump().items() if v is not None}
    task = service.tasks.update(task_id, **changes)
    if task is None:
        raise HTTPException(status_code=404, detail="Task not found")
    # Log status transitions to activity feed
    try:
        action = "task.completed" if task.status == "done" else "task.updated"
        service.activity.log(
            action=action,
            entity_type="task",
            entity_id=task.task_id,
            entity_title=task.title,
            status=task.status,
        )
    except Exception:
        pass
    return task.to_dict()


@app.delete("/tasks/{task_id}", status_code=204)
def delete_task(task_id: str) -> None:
    if not service.tasks.delete(task_id):
        raise HTTPException(status_code=404, detail="Task not found")


# ── Notifications ──────────────────────────────────────────────────────────────

class NotificationCreate(BaseModel):
    recipient_id: str
    title: str
    body: str = ""
    type: str = "general"
    entity_type: Optional[str] = None
    entity_id: Optional[str] = None


@app.get("/notifications")
def list_notifications(
    recipient_id: str = "user-1",
    unread_only: bool = False,
    limit: int = 50,
) -> list[dict]:
    return [n.to_dict() for n in service.notifications.list(recipient_id=recipient_id, unread_only=unread_only, limit=limit)]


@app.get("/notifications/unread-count")
def unread_notification_count(recipient_id: str = "user-1") -> dict:
    return {"count": service.notifications.count_unread(recipient_id)}


@app.post("/notifications", status_code=201)
def create_notification(payload: NotificationCreate) -> dict:
    notif = service.notifications.create(
        recipient_id=payload.recipient_id,
        title=payload.title,
        body=payload.body,
        type=payload.type,
        entity_type=payload.entity_type,
        entity_id=payload.entity_id,
    )
    return notif.to_dict()


@app.post("/notifications/{notification_id}/read", status_code=204)
def mark_notification_read(notification_id: str) -> None:
    service.notifications.mark_read(notification_id)


@app.post("/notifications/read-all", status_code=204)
def mark_all_notifications_read(recipient_id: str = "user-1") -> None:
    service.notifications.mark_all_read(recipient_id)


@app.delete("/notifications/{notification_id}", status_code=204)
def delete_notification(notification_id: str) -> None:
    service.notifications.delete(notification_id)


# ── Activity Log ───────────────────────────────────────────────────────────────


@app.get("/activity")
def list_activity(
    org_id: str = "org-1",
    entity_type: Optional[str] = None,
    action_prefix: Optional[str] = None,
    limit: int = 50,
) -> list[dict]:
    entries = service.activity.list_for_org(
        org_id=org_id,
        limit=limit,
        entity_type=entity_type,
        action_prefix=action_prefix,
    )
    return [e.to_dict() for e in entries]


@app.get("/activity/{entity_type}/{entity_id}")
def list_activity_for_entity(
    entity_type: str,
    entity_id: str,
    limit: int = 50,
) -> list[dict]:
    entries = service.activity.list_for_entity(
        entity_type=entity_type,
        entity_id=entity_id,
        limit=limit,
    )
    return [e.to_dict() for e in entries]


@app.get("/agents")
def list_agents() -> dict:
    return {"agents": [asdict(manifest) for manifest in service.registry.list_active()]}


@app.post("/agents/scaffold")
def scaffold_agent(payload: ScaffoldAgentPayload) -> dict:
    from scripts.create_agent_from_template import create_agent_from_template

    paths = create_agent_from_template(
        agent_id=payload.id,
        name=payload.name,
        purpose=payload.purpose,
        owner=payload.owner,
    )
    return {"status": "created", "paths": paths}


@app.post("/evals/run")
def run_evals(payload: EvalPayload) -> dict:
    suite = payload.suite or "core-routing"
    try:
        report = service.eval_harness.run_suite(suite, service.manager)
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    return report.to_dict()


@app.get("/admin/dashboard")
def admin_dashboard(request: Request) -> dict:
    admin_auth.require(request)
    return service.get_dashboard_metrics()


@app.get("/admin/tools")
def admin_tools(request: Request) -> dict:
    admin_auth.require(request)
    return {"tools": service.tools.list_tools()}


@app.post("/admin/tools/mcp/register")
def admin_mcp_register(payload: dict, request: Request) -> dict:
    admin_auth.require(request)
    required = ["server_id", "name", "endpoint"]
    missing = [field for field in required if not payload.get(field)]
    if missing:
        raise HTTPException(status_code=400, detail=f"missing fields: {', '.join(missing)}")

    from packages.tools.mcp import MCPServer

    try:
        server = service.mcp.register(
            MCPServer(
                server_id=str(payload["server_id"]),
                name=str(payload["name"]),
                endpoint=str(payload["endpoint"]),
                auth_type=str(payload.get("auth_type") or "none"),
                enabled=bool(payload.get("enabled", True)),
            )
        )
    except ValueError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    return {"server": asdict(server)}


@app.post("/admin/tools/mcp/{server_id}/enable")
def admin_mcp_enable(server_id: str, payload: dict, request: Request) -> dict:
    admin_auth.require(request)
    enabled = bool(payload.get("enabled", True))
    try:
        server = service.mcp.set_enabled(server_id, enabled=enabled)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="server not found") from exc
    return {"server": asdict(server)}


@app.post("/admin/agents/{agent_id}/status")
def admin_agent_status(agent_id: str, payload: dict, request: Request) -> dict:
    admin_auth.require(request)
    status = str(payload.get("status") or "").strip()
    if not status:
        raise HTTPException(status_code=400, detail="status is required")
    try:
        manifest = service.registry.update_status(agent_id, status)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="agent not found") from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return {"agent": asdict(manifest)}


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)) -> dict:
    """Extract text from an uploaded file and return a context_id.

    Supported types: .txt, .md, .csv, .pdf (requires pypdf), .docx (requires python-docx),
    .png/.jpg/.jpeg/.webp (requires openai — routed to GPT-4o vision).

    Include the returned context_id in the ``context_ids`` field of a /chat request
    to inject the extracted text as specialist context.
    """
    import csv as _csv
    import io
    import uuid

    if not file.filename:
        raise HTTPException(status_code=400, detail="filename is required")

    raw = await file.read()
    name = file.filename.lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""

    text = ""
    doc_type = ext

    if ext in ("txt", "md"):
        try:
            text = raw.decode("utf-8", errors="replace")
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Could not decode text file: {exc}") from exc

    elif ext == "csv":
        try:
            reader = _csv.reader(io.StringIO(raw.decode("utf-8", errors="replace")))
            rows = list(reader)
            text = "\n".join(", ".join(row) for row in rows[:500])  # cap at 500 rows
            if len(rows) > 500:
                text += f"\n... ({len(rows) - 500} rows truncated)"
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"CSV parse error: {exc}") from exc

    elif ext == "pdf":
        try:
            import pypdf  # type: ignore
            reader = pypdf.PdfReader(io.BytesIO(raw))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(p for p in pages if p.strip())
        except ImportError:
            raise HTTPException(
                status_code=422,
                detail="PDF extraction requires 'pypdf'. Install it: pip install pypdf",
            )
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"PDF parse error: {exc}") from exc

    elif ext == "docx":
        try:
            import docx  # type: ignore
            doc = docx.Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise HTTPException(
                status_code=422,
                detail="DOCX extraction requires 'python-docx'. Install it: pip install python-docx",
            )
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"DOCX parse error: {exc}") from exc

    elif ext in ("png", "jpg", "jpeg", "webp", "gif"):
        # Route to GPT-4o vision
        try:
            import base64
            from openai import OpenAI  # type: ignore
            b64 = base64.b64encode(raw).decode("utf-8")
            mime = f"image/{ext if ext != 'jpg' else 'jpeg'}"
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))
            resp = client.chat.completions.create(
                model="gpt-4o",
                max_tokens=1000,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                        {"type": "text", "text": "Describe this image in detail. Extract all text, tables, charts, and key information visible."},
                    ],
                }],
            )
            text = resp.choices[0].message.content or ""
            doc_type = "image_vision"
        except ImportError:
            raise HTTPException(status_code=422, detail="Image analysis requires the 'openai' package")
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Vision analysis error: {exc}") from exc

    else:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type: .{ext}. Supported: txt, md, csv, pdf, docx, png, jpg, jpeg, webp",
        )

    if not text.strip():
        raise HTTPException(status_code=422, detail="No text could be extracted from the file")

    context_id = f"ctx_{uuid.uuid4().hex[:12]}"
    _UPLOAD_STORE[context_id] = {
        "filename": file.filename,
        "text": text,
        "type": doc_type,
        "chars": len(text),
    }

    return {
        "context_id": context_id,
        "filename": file.filename,
        "type": doc_type,
        "text_length": len(text),
    }


# ── Process mapping endpoints ─────────────────────────────────────────────────

@app.get("/processes")
def list_processes(org_id: str = "org-1") -> list[dict]:
    docs = service.processes.list(org_id=org_id)
    return [d.to_dict() for d in docs]


@app.post("/processes", status_code=201)
def create_process(payload: ProcessCreatePayload) -> dict:
    from packages.common.models import ProcessDocument, ProcessStep
    steps = [ProcessStep(**s) if isinstance(s, dict) else s for s in payload.steps]
    doc = ProcessDocument(
        id="",
        org_id=payload.org_id,
        process_name=payload.process_name,
        trigger=payload.trigger,
        steps=steps,
        decision_points=payload.decision_points,
        roles=payload.roles,
        tools=payload.tools,
        exceptions=payload.exceptions,
        kpis=payload.kpis,
        mermaid_flowchart=payload.mermaid_flowchart,
        mermaid_swimlane=payload.mermaid_swimlane,
        completeness_score=0.0,
        status=payload.status,
    )
    created = service.processes.create(doc)
    return created.to_dict()


@app.get("/processes/analytics")
def process_analytics(org_id: str = "org-1") -> dict:
    return service.process_analytics.org_health(org_id=org_id)


@app.get("/processes/{process_id}")
def get_process(process_id: str) -> dict:
    doc = service.processes.get(process_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Process not found")
    return doc.to_dict()


@app.put("/processes/{process_id}")
def update_process(process_id: str, payload: ProcessUpdatePayload) -> dict:
    changes: dict = {}
    for field in ("process_name", "trigger", "steps", "decision_points",
                  "roles", "tools", "exceptions", "kpis",
                  "mermaid_flowchart", "mermaid_swimlane", "status"):
        val = getattr(payload, field, None)
        if val is not None:
            changes[field] = val
    if not changes:
        raise HTTPException(status_code=400, detail="No changes provided")
    try:
        result = service.processes.update(
            process_id,
            changes=changes,
            bump=payload.bump,
            author=payload.author,
            changelog_entry=payload.changelog_entry,
        )
        # Major bump held for approval — return 202 Accepted with pending payload
        if isinstance(result, dict) and result.get("status") == "pending_approval":
            return JSONResponse(status_code=202, content=result)
        return result.to_dict()
    except KeyError:
        raise HTTPException(status_code=404, detail="Process not found")


@app.get("/processes/{process_id}/history")
def process_history(process_id: str) -> list[dict]:
    doc = service.processes.get(process_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Process not found")
    return service.processes.history(process_id)


@app.get("/processes/{process_id}/versions/{version}")
def get_process_version(process_id: str, version: str) -> dict:
    doc = service.processes.get_version(process_id, version)
    if doc is None:
        raise HTTPException(status_code=404, detail="Version not found")
    return doc.to_dict()


@app.delete("/processes/{process_id}", status_code=204)
def delete_process(process_id: str) -> None:
    doc = service.processes.get(process_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Process not found")
    service.processes.delete(process_id)


@app.get("/processes/{process_id}/completeness")
def process_completeness(process_id: str) -> dict:
    doc = service.processes.get(process_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Process not found")
    return service.processes.completeness_breakdown(doc)


# ─────────────────────────────────────────────────────────────────────────────

# ── File storage endpoints ────────────────────────────────────────────────────

@app.get("/files")
def list_files(org_id: str = "org-1") -> list[dict]:
    return [f.to_dict() for f in service.storage.list_files(org_id=org_id)]


@app.get("/files/{file_id}")
def download_file(file_id: str):
    try:
        meta, content = service.storage.retrieve(file_id)
    except (KeyError, FileNotFoundError) as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    return FileResponse(
        path=meta.storage_path,
        filename=meta.filename,
        media_type=meta.mime_type,
    )


@app.delete("/files/{file_id}", status_code=204)
def delete_file(file_id: str) -> None:
    service.storage.delete(file_id)


# ── Document generation endpoints ────────────────────────────────────────────


class DocGenPayload(BaseModel):
    title: str
    document_type: str = "memo"
    format: str = "docx"  # docx, pptx, xlsx, pdf
    sections: list[dict] = []
    metadata: dict[str, Any] = {}
    org_id: str = "org-1"


@app.post("/documents/generate")
def generate_document(payload: DocGenPayload) -> dict:
    if service.docgen is None:
        raise HTTPException(status_code=501, detail="Document generation not available. Install python-docx, python-pptx, openpyxl.")
    from packages.docgen.generators.base import DocumentContent, DocumentSection
    sections = [DocumentSection(**s) for s in payload.sections]
    content = DocumentContent(
        title=payload.title,
        document_type=payload.document_type,
        sections=sections,
        metadata=payload.metadata,
    )
    stored = service.docgen.generate(content, format=payload.format, org_id=payload.org_id)
    return stored.to_dict()


@app.get("/documents")
def list_documents(org_id: str = "org-1", workspace_id: Optional[str] = None, q: Optional[str] = None) -> list[dict]:
    files = service.storage.list_files(org_id=org_id)
    results = [f.to_dict() for f in files if f.metadata.get("format") in ("docx", "pptx", "xlsx", "pdf")]
    if workspace_id:
        results = [r for r in results if r.get("metadata", {}).get("workspace_id") == workspace_id]
    if q:
        q_lower = q.strip().lower()
        results = [r for r in results if q_lower in r.get("filename", "").lower()]
    return results


# ── Template endpoints ───────────────────────────────────────────────────────

@app.get("/templates")
def list_templates(org_id: str = "org-1", category: Optional[str] = None) -> list[dict]:
    return [t.to_dict() for t in service.templates.list_templates(org_id=org_id, category=category)]


@app.get("/templates/{template_id}")
def get_template(template_id: str) -> dict:
    t = service.templates.get(template_id)
    if t is None:
        raise HTTPException(status_code=404, detail="Template not found")
    return t.to_dict()


# ── Credential / Integration endpoints ───────────────────────────────────────

@app.get("/credentials")
def list_credentials(org_id: str = "org-1") -> list[dict]:
    return [c.to_dict() for c in service.credentials.list_credentials(org_id=org_id)]


@app.get("/integrations/status")
def integration_status(org_id: str = "org-1") -> dict:
    """Check which integrations are connected."""
    providers = ["google", "slack", "jira", "linear", "confluence", "notion", "salesforce", "hubspot", "gmail", "outlook"]
    return {p: service.credentials.has_credential(p, org_id) for p in providers}


# ── KPI / Analytics endpoints ────────────────────────────────────────────────


class KPICreatePayload(BaseModel):
    name: str
    unit: str
    target_value: Optional[float] = None  # frontend uses target_value
    target: Optional[float] = None        # legacy alias
    direction: str = "higher_is_better"
    category: str = ""
    frequency: str = "monthly"
    data_source: str = "manual"
    org_id: str = "org-1"


class KPIDataPayload(BaseModel):
    value: float
    source: str = "manual"


def _normalize_kpi(raw: dict) -> dict:
    """Convert backend KPI shape → frontend KPI shape."""
    latest = raw.get("latest_value") or 0.0
    target = raw.get("target") or raw.get("target_value") or 0.0
    on_target = raw.get("on_target")
    direction = raw.get("direction", "higher_is_better")

    if latest is None or latest == 0.0 and target == 0.0:
        status = "at_risk"
    elif direction == "higher_is_better":
        pct = (latest / target) if target else 0
        status = "on_track" if pct >= 0.9 else ("at_risk" if pct >= 0.7 else "behind")
    else:
        pct = (latest / target) if target else 0
        status = "on_track" if pct <= 1.1 else ("at_risk" if pct <= 1.3 else "behind")

    return {
        "kpi_id": raw.get("kpi_id", ""),
        "name": raw.get("name", ""),
        "unit": raw.get("unit", ""),
        "current_value": latest if latest is not None else 0.0,
        "target_value": target,
        "direction": direction,
        "category": raw.get("category", ""),
        "status": status,
    }


@app.get("/kpis")
def list_kpis(org_id: str = "org-1") -> list[dict]:
    return [_normalize_kpi(k) for k in service.kpis.kpi_status(org_id=org_id)]


@app.post("/kpis", status_code=201)
def create_kpi(payload: KPICreatePayload) -> dict:
    target = payload.target_value if payload.target_value is not None else payload.target
    kpi = service.kpis.create_kpi(
        name=payload.name, unit=payload.unit, target=target,
        frequency=payload.frequency, data_source=payload.data_source, org_id=payload.org_id)
    return _normalize_kpi({
        **kpi.to_dict(),
        "latest_value": 0.0,
        "on_target": None,
        "direction": payload.direction,
        "category": payload.category,
    })


@app.post("/kpis/{kpi_id}/data", status_code=201)
def record_kpi_data(kpi_id: str, payload: KPIDataPayload) -> dict:
    dp = service.kpis.record_data_point(kpi_id, value=payload.value, source=payload.source)
    # Auto-sync any KRs linked to this KPI
    try:
        linked_krs = service.okrs.list_key_results_by_kpi(kpi_id)
        for kr in linked_krs:
            service.okrs.sync_kr_from_kpi_value(kr.kr_id, payload.value)
    except Exception:
        pass  # KPI sync is best-effort; never fail the data-record call
    return dp.to_dict()


@app.get("/kpis/{kpi_id}/trend")
def kpi_trend(kpi_id: str, limit: int = 30) -> list[dict]:
    return [dp.to_dict() for dp in service.kpis.get_trend(kpi_id, limit=limit)]


# ── OKR endpoints ────────────────────────────────────────────────────────────


class ObjectivePayload(BaseModel):
    name: str
    description: str = ""
    period: str = "Q1 2026"
    org_id: str = "org-1"


class KeyResultPayload(BaseModel):
    name: str
    target: float
    unit: str = "%"


class KRProgressPayload(BaseModel):
    current: float


# ── Full OKR Pydantic models ──────────────────────────────────────────────────

class ObjectiveCreate(BaseModel):
    org_id: str = "default"
    workspace_id: Optional[str] = None
    title: str
    description: str = ""
    owner: str = "user"
    collaborators: list[str] = []
    parent_id: Optional[str] = None
    level: str = "team"
    period: str = "2026-Q1"
    status: str = "active"
    confidence: float = 0.7
    rationale: str = ""
    linked_initiatives: list[str] = []
    linked_docs: list[str] = []


class KeyResultCreate(BaseModel):
    title: str
    metric_type: str = "number"
    baseline: float = 0.0
    current_value: float = 0.0
    target_value: float = 100.0
    unit: str = ""
    owner: str = "user"
    data_source: str = ""
    update_cadence: str = "weekly"
    status: str = "active"
    confidence: float = 0.7
    due_date: str = ""
    notes: str = ""


class InitiativeCreate(BaseModel):
    title: str
    owner: str = "user"
    kr_id: Optional[str] = None
    status: str = "not_started"
    due_date: str = ""
    description: str = ""
    org_id: str = "default"
    workspace_id: Optional[str] = None


class CheckInCreate(BaseModel):
    author: str = "user"
    status: str = "on_track"
    confidence: float = 0.7
    highlights: str = ""
    blockers: str = ""
    next_steps: str = ""


# ── Workspace Pydantic models ─────────────────────────────────────────────────

class WorkspaceCreate(BaseModel):
    name: str
    description: str = ""
    icon: str = "🗂️"
    color: str = "#6366f1"
    type: str = "team"
    owner: str = "user"
    org_id: str = "default"
    visibility: str = "open"
    default_view: str = "chat"


class WorkspaceMemberAdd(BaseModel):
    user_id: str
    role: str = "editor"


class WorkspaceLinkCreate(BaseModel):
    entity_type: str
    entity_id: str


@app.get("/okrs")
def list_okrs(
    org_id: str = "default",
    workspace_id: Optional[str] = None,
    level: Optional[str] = None,
    period: Optional[str] = None,
    status: Optional[str] = None,
) -> list[dict]:
    objs = service.okrs.list_objectives(
        org_id=org_id,
        workspace_id=workspace_id,
        parent_id=None,  # return all, not just root
        level=level,
        period=period,
        status=status,
    )
    return [asdict(o) for o in objs]


@app.post("/okrs", status_code=201)
def create_objective_full(payload: ObjectiveCreate) -> dict:
    obj = service.okrs.create_objective(
        title=payload.title,
        period=payload.period,
        org_id=payload.org_id,
        description=payload.description,
        owner=payload.owner,
        level=payload.level,
        parent_id=payload.parent_id,
        workspace_id=payload.workspace_id,
        collaborators=payload.collaborators,
        rationale=payload.rationale,
        confidence=payload.confidence,
    )
    return asdict(obj)


@app.get("/okrs/{obj_id}")
def get_objective(obj_id: str) -> dict:
    result = service.okrs.get_objective_with_details(obj_id)
    if not result:
        raise HTTPException(status_code=404, detail="Not found")
    return result


@app.put("/okrs/{obj_id}")
def update_objective(obj_id: str, payload: dict) -> dict:
    obj = service.okrs.update_objective(obj_id, **payload)
    if obj is None:
        raise HTTPException(status_code=404, detail="Not found")
    return asdict(obj)


@app.delete("/okrs/{obj_id}", status_code=204)
def delete_objective(obj_id: str) -> None:
    service.okrs.update_objective(obj_id, status="cancelled")


@app.get("/okrs/{obj_id}/children")
def list_objective_children(obj_id: str, org_id: str = "org-1") -> list[dict]:
    children = service.okrs.get_children(obj_id, org_id=org_id)
    return [asdict(c) for c in children]


@app.post("/okrs/{obj_id}/key-results", status_code=201)
def create_key_result_full(obj_id: str, payload: KeyResultCreate) -> dict:
    obj = service.okrs.get_objective(obj_id)
    if obj is None:
        raise HTTPException(status_code=404, detail="Not found")
    kr = service.okrs.create_key_result(
        objective_id=obj_id,
        title=payload.title,
        target_value=payload.target_value,
        org_id=obj.org_id,
        unit=payload.unit,
        baseline=payload.baseline,
        owner=payload.owner,
        metric_type=payload.metric_type,
        data_source=payload.data_source,
        update_cadence=payload.update_cadence,
        due_date=payload.due_date,
        confidence=payload.confidence,
    )
    return asdict(kr)


@app.put("/okrs/key-results/{kr_id}")
def update_key_result_full(kr_id: str, payload: dict) -> dict:
    kr = service.okrs.update_key_result(kr_id, **payload)
    if kr is None:
        raise HTTPException(status_code=404, detail="Not found")
    return asdict(kr)


@app.post("/okrs/{obj_id}/initiatives", status_code=201)
def create_initiative(obj_id: str, payload: InitiativeCreate) -> dict:
    init = service.okrs.create_initiative(
        title=payload.title,
        objective_id=obj_id,
        org_id=payload.org_id,
        owner=payload.owner,
        kr_id=payload.kr_id,
        description=payload.description,
        due_date=payload.due_date,
        workspace_id=payload.workspace_id,
    )
    return asdict(init)


@app.post("/okrs/{obj_id}/check-in", status_code=201)
def okr_checkin(obj_id: str, payload: CheckInCreate) -> dict:
    obj = service.okrs.get_objective(obj_id)
    if obj is None:
        raise HTTPException(status_code=404, detail="Not found")
    checkin = service.okrs.add_checkin(
        objective_id=obj_id,
        notes="",
        confidence=payload.confidence,
        status=payload.status,
        author=payload.author,
        org_id=obj.org_id,
        highlights=payload.highlights,
        blockers=payload.blockers,
        next_steps=payload.next_steps,
    )
    # Log to activity feed
    try:
        service.activity.log(
            action="okr.checkin",
            entity_type="objective",
            entity_id=obj_id,
            entity_title=obj.title,
            actor_id=payload.author or "system",
            org_id=obj.org_id,
            status=payload.status,
            confidence=payload.confidence,
        )
    except Exception:
        pass
    return asdict(checkin)


class InitiativeUpdate(BaseModel):
    status: Optional[str] = None
    owner: Optional[str] = None
    title: Optional[str] = None
    due_date: Optional[str] = None
    description: Optional[str] = None


@app.put("/okrs/initiatives/{initiative_id}", status_code=200)
def update_initiative(initiative_id: str, payload: InitiativeUpdate) -> dict:
    service.okrs.update_initiative(
        initiative_id,
        status=payload.status,
        owner=payload.owner,
        title=payload.title,
        due_date=payload.due_date,
        description=payload.description,
    )
    return {"ok": True}


@app.put("/okrs/key-results/{kr_id}/progress")
def update_kr_progress(kr_id: str, payload: KRProgressPayload) -> dict:
    return service.okrs.update_key_result_progress(kr_id, current=payload.current)


class LinkKPIPayload(BaseModel):
    kpi_id: Optional[str] = None   # None = unlink


@app.post("/okrs/key-results/{kr_id}/link-kpi", status_code=200)
def link_kr_to_kpi(kr_id: str, payload: LinkKPIPayload) -> dict:
    """Associate a Key Result with a KPI for automatic value sync.
    Pass kpi_id=null to remove the link."""
    ok = service.okrs.link_kr_to_kpi(kr_id, payload.kpi_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Key result not found")
    # If linking (not unlinking), immediately sync the latest KPI value
    if payload.kpi_id:
        try:
            latest = service.kpis.get_latest_value(payload.kpi_id)
            if latest is not None:
                service.okrs.sync_kr_from_kpi_value(kr_id, latest)
        except Exception:
            pass
    return {"ok": True, "kr_id": kr_id, "kpi_id": payload.kpi_id}


@app.get("/okrs/{obj_id}/hierarchy")
def get_okr_hierarchy(obj_id: str) -> dict:
    """Return a full nested tree: objective + all descendant children + their KRs."""
    tree = service.okrs.get_hierarchy(obj_id)
    if not tree:
        raise HTTPException(status_code=404, detail="Objective not found")
    return tree


# ── Workspace endpoints ───────────────────────────────────────────────────────

@app.get("/workspaces")
def list_workspaces(org_id: str = "default") -> list[dict]:
    workspaces = service.workspaces.list(org_id=org_id)
    return [asdict(w) for w in workspaces]


@app.post("/workspaces", status_code=201)
def create_workspace(payload: WorkspaceCreate) -> dict:
    ws = service.workspaces.create(
        name=payload.name,
        org_id=payload.org_id,
        owner=payload.owner,
        type=payload.type,
        description=payload.description,
        icon=payload.icon,
        color=payload.color,
        visibility=payload.visibility,
        default_view=payload.default_view,
    )
    return asdict(ws)


@app.get("/workspaces/{workspace_id}")
def get_workspace(workspace_id: str) -> dict:
    ws = service.workspaces.get(workspace_id)
    if ws is None:
        raise HTTPException(status_code=404, detail="Not found")
    return asdict(ws)


@app.put("/workspaces/{workspace_id}")
def update_workspace(workspace_id: str, payload: dict) -> dict:
    ws = service.workspaces.update(workspace_id, **payload)
    if ws is None:
        raise HTTPException(status_code=404, detail="Not found")
    return asdict(ws)


@app.post("/workspaces/{workspace_id}/archive")
def archive_workspace(workspace_id: str) -> dict:
    service.workspaces.archive(workspace_id)
    return {"status": "archived", "workspace_id": workspace_id}


@app.delete("/workspaces/{workspace_id}", status_code=204)
def delete_workspace(workspace_id: str) -> None:
    """Delete (archive) a workspace. Admin-only by convention; enforced in the frontend."""
    service.workspaces.archive(workspace_id)


@app.get("/workspaces/{workspace_id}/members")
def list_workspace_members(workspace_id: str) -> list[dict]:
    members = service.workspaces.list_members(workspace_id)
    return [asdict(m) for m in members]


@app.post("/workspaces/{workspace_id}/members", status_code=201)
def add_workspace_member(workspace_id: str, payload: WorkspaceMemberAdd) -> dict:
    member = service.workspaces.add_member(workspace_id, user_id=payload.user_id, role=payload.role)
    return asdict(member)


@app.delete("/workspaces/{workspace_id}/members/{user_id}", status_code=204)
def remove_workspace_member(workspace_id: str, user_id: str) -> None:
    service.workspaces.remove_member(workspace_id, user_id)


@app.post("/workspaces/{workspace_id}/link", status_code=201)
def link_workspace_entity(workspace_id: str, payload: WorkspaceLinkCreate) -> dict:
    link = service.workspaces.link_entity(workspace_id, entity_type=payload.entity_type, entity_id=payload.entity_id)
    return asdict(link)


@app.get("/workspaces/{workspace_id}/overview")
def workspace_overview(workspace_id: str) -> dict:
    ws = service.workspaces.get(workspace_id)
    if ws is None:
        raise HTTPException(status_code=404, detail="Not found")
    members = service.workspaces.list_members(workspace_id)
    return {
        "workspace": asdict(ws),
        "member_count": len(members),
        "recent_okrs": [],
    }


# ── Project endpoints ────────────────────────────────────────────────────────


class ProjectCreate(BaseModel):
    name: str
    description: str = ""
    color: str = "#6366f1"
    icon: str = "📁"


class ProjectUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    color: Optional[str] = None
    icon: Optional[str] = None


@app.get("/workspaces/{workspace_id}/projects")
def list_projects(workspace_id: str) -> list[dict]:
    return [p.to_dict() for p in service.projects.list(workspace_id)]


@app.post("/workspaces/{workspace_id}/projects", status_code=201)
def create_project(workspace_id: str, payload: ProjectCreate) -> dict:
    project = service.projects.create(
        workspace_id=workspace_id,
        name=payload.name,
        description=payload.description,
        color=payload.color,
        icon=payload.icon,
    )
    return project.to_dict()


@app.put("/projects/{project_id}")
def update_project(project_id: str, payload: ProjectUpdate) -> dict:
    project = service.projects.update(project_id, **payload.model_dump(exclude_none=True))
    if project is None:
        raise HTTPException(status_code=404, detail="Not found")
    return project.to_dict()


@app.delete("/projects/{project_id}", status_code=204)
def delete_project(project_id: str) -> None:
    service.projects.delete(project_id)


# ── Finance endpoints ────────────────────────────────────────────────────────


class InvoiceItemPayload(BaseModel):
    description: str
    quantity: float = 1.0
    unit_price: float = 0.0


class InvoiceCreatePayload(BaseModel):
    client_name: str
    client_address: str = ""
    items: list[InvoiceItemPayload] = []
    tax_rate: float = 0.0
    due_date: str = ""
    org_id: str = "org-1"


class BudgetCategoryPayload(BaseModel):
    name: str
    planned_amount: float
    period: str = "monthly"
    org_id: str = "org-1"


class ExpensePayload(BaseModel):
    category_id: str
    amount: float
    description: str = ""


@app.get("/invoices")
def list_invoices(org_id: str = "org-1") -> list[dict]:
    return [i.to_dict() for i in service.invoices.list_invoices(org_id=org_id)]


@app.post("/invoices", status_code=201)
def create_invoice(payload: InvoiceCreatePayload) -> dict:
    from packages.finance.invoice_service import InvoiceItem
    items = [InvoiceItem(description=i.description, quantity=i.quantity, unit_price=i.unit_price)
             for i in payload.items]
    inv = service.invoices.create_invoice(
        client_name=payload.client_name, client_address=payload.client_address,
        items=items, tax_rate=payload.tax_rate, due_date=payload.due_date, org_id=payload.org_id)
    return inv.to_dict()


@app.get("/budget/status")
def budget_status(org_id: str = "org-1") -> list[dict]:
    return service.budgets.budget_status(org_id=org_id)


@app.post("/budget/categories", status_code=201)
def create_budget_category(payload: BudgetCategoryPayload) -> dict:
    cat = service.budgets.create_category(
        name=payload.name, planned_amount=payload.planned_amount,
        period=payload.period, org_id=payload.org_id)
    return cat.to_dict()


@app.post("/budget/expenses", status_code=201)
def record_expense(payload: ExpensePayload) -> dict:
    exp = service.budgets.record_expense(
        category_id=payload.category_id, amount=payload.amount, description=payload.description)
    return exp.to_dict()


# ── Brand endpoints ──────────────────────────────────────────────────────────

@app.get("/brand")
def get_brand(org_id: str = "org-1") -> dict:
    return service.brand.get_brand_or_default(org_id=org_id)


class BrandUpdatePayload(BaseModel):
    primary_color: Optional[str] = None
    secondary_color: Optional[str] = None
    accent_color: Optional[str] = None
    font_family: Optional[str] = None
    company_name: Optional[str] = None
    tagline: Optional[str] = None
    voice_tone: Optional[str] = None


@app.put("/brand")
def update_brand(payload: BrandUpdatePayload, org_id: str = "org-1") -> dict:
    changes = {k: v for k, v in payload.model_dump().items() if v is not None}
    return service.brand.update_brand(org_id=org_id, changes=changes)


# ── Events endpoint ──────────────────────────────────────────────────────────

@app.get("/events")
def list_events(limit: int = 50) -> list[dict]:
    from dataclasses import asdict
    return [asdict(e) for e in service.events.recent_events(limit=limit)]


# ── Code Interpreter endpoints ───────────────────────────────────────────────

class CodeRunPayload(BaseModel):
    code: str
    data_files: dict[str, str] = {}
    org_id: str = "org-1"


@app.post("/analysis/run")
def run_code(payload: CodeRunPayload) -> dict:
    return service.interpreter.run(payload.code, data_files=payload.data_files, org_id=payload.org_id)


# ── Meeting endpoints ─────────────────────────────────────────────────────────

class MeetingCreatePayload(BaseModel):
    title: str
    scheduled_at: str
    attendees: list[str] = []
    agenda: list[str] = []
    duration_minutes: int = 60
    org_id: str = "org-1"


class MeetingNotesPayload(BaseModel):
    raw_text: str
    org_id: str = "org-1"


@app.get("/meetings")
def list_meetings(org_id: str = "org-1", status: Optional[str] = None) -> list[dict]:
    from dataclasses import asdict
    return [asdict(m) for m in service.meetings.list_meetings(org_id=org_id, status=status)]


@app.post("/meetings", status_code=201)
def create_meeting(payload: MeetingCreatePayload) -> dict:
    from dataclasses import asdict
    m = service.meetings.create_meeting(
        title=payload.title, scheduled_at=payload.scheduled_at,
        attendees=payload.attendees, agenda=payload.agenda,
        duration_minutes=payload.duration_minutes, org_id=payload.org_id)
    return asdict(m)


@app.post("/meetings/{meeting_id}/notes", status_code=201)
def process_meeting_notes(meeting_id: str, payload: MeetingNotesPayload) -> dict:
    from dataclasses import asdict
    note = service.meetings.process_notes(meeting_id, payload.raw_text, org_id=payload.org_id)
    result = asdict(note)
    result["action_items"] = [asdict(a) for a in note.action_items]
    return result


@app.get("/meetings/action-items")
def list_action_items(org_id: str = "org-1", status: str = "open", owner: Optional[str] = None) -> list[dict]:
    from dataclasses import asdict
    return [asdict(a) for a in service.meetings.list_action_items(org_id=org_id, status=status, owner=owner)]


@app.post("/meetings/action-items/{item_id}/complete")
def complete_action_item(item_id: str) -> dict:
    service.meetings.complete_action_item(item_id)
    return {"status": "done"}


# ── Org context endpoints ─────────────────────────────────────────────────────

class OrgProfilePayload(BaseModel):
    company_name: str
    industry: str = ""
    stage: str = "growth"
    fiscal_year_end: str = "12-31"
    headcount: int = 0
    founded_year: Optional[int] = None
    mission: str = ""
    values: list[str] = []
    org_id: str = "org-1"


class PersonPayload(BaseModel):
    name: str
    role: str
    department: str
    email: str = ""
    reports_to: str = ""
    org_id: str = "org-1"


class PriorityPayload(BaseModel):
    title: str
    description: str
    owner: str
    due_date: str
    tags: list[str] = []
    org_id: str = "org-1"


@app.get("/org/profile")
def get_org_profile(org_id: str = "org-1") -> dict:
    from dataclasses import asdict
    p = service.org_context.get_profile(org_id)
    return asdict(p) if p else {}


@app.put("/org/profile")
def upsert_org_profile(payload: OrgProfilePayload) -> dict:
    from dataclasses import asdict
    from packages.org_context import OrgProfile
    profile = OrgProfile(
        org_id=payload.org_id, company_name=payload.company_name,
        industry=payload.industry, stage=payload.stage,
        fiscal_year_end=payload.fiscal_year_end, headcount=payload.headcount,
        founded_year=payload.founded_year, mission=payload.mission,
        values=payload.values)
    return asdict(service.org_context.upsert_profile(profile))


@app.get("/org/people")
def list_people(org_id: str = "org-1", department: Optional[str] = None) -> list[dict]:
    from dataclasses import asdict
    return [asdict(p) for p in service.org_context.list_people(org_id, department)]


@app.post("/org/people", status_code=201)
def add_person(payload: PersonPayload) -> dict:
    from dataclasses import asdict
    from packages.org_context import Person
    from uuid import uuid4
    p = Person(person_id=f"person_{uuid4().hex[:12]}", name=payload.name,
                role=payload.role, department=payload.department,
                email=payload.email, reports_to=payload.reports_to, org_id=payload.org_id)
    return asdict(service.org_context.upsert_person(p))


@app.get("/org/chart")
def org_chart(org_id: str = "org-1") -> dict:
    return service.org_context.org_chart(org_id)


@app.get("/org/priorities")
def list_priorities(org_id: str = "org-1") -> list[dict]:
    from dataclasses import asdict
    return [asdict(p) for p in service.org_context.list_priorities(org_id)]


@app.post("/org/priorities", status_code=201)
def add_priority(payload: PriorityPayload) -> dict:
    from dataclasses import asdict
    return asdict(service.org_context.add_priority(
        title=payload.title, description=payload.description,
        owner=payload.owner, due_date=payload.due_date,
        org_id=payload.org_id, tags=payload.tags))


# ── Decision log endpoints ────────────────────────────────────────────────────

class DecisionPayload(BaseModel):
    title: str
    context: str
    rationale: str
    owner: str
    options_considered: list[str] = []
    tags: list[str] = []
    reversibility: str = "reversible"
    confidence: float = 0.8
    related_run_id: str = ""
    org_id: str = "org-1"


@app.get("/decisions")
def list_decisions(org_id: str = "org-1", tag: Optional[str] = None) -> list[dict]:
    from dataclasses import asdict
    return [asdict(d) for d in service.decisions.list_decisions(org_id, tag=tag)]


@app.post("/decisions", status_code=201)
def log_decision(payload: DecisionPayload) -> dict:
    from dataclasses import asdict
    return asdict(service.decisions.log(
        title=payload.title, context=payload.context, rationale=payload.rationale,
        owner=payload.owner, options_considered=payload.options_considered,
        org_id=payload.org_id, tags=payload.tags, reversibility=payload.reversibility,
        confidence=payload.confidence, related_run_id=payload.related_run_id))


@app.get("/decisions/search")
def search_decisions(q: str, org_id: str = "org-1") -> list[dict]:
    from dataclasses import asdict
    return [asdict(d) for d in service.decisions.search(q, org_id)]


@app.post("/decisions/{decision_id}/outcome")
def record_decision_outcome(decision_id: str, outcome: str) -> dict:
    service.decisions.record_outcome(decision_id, outcome)
    return {"status": "recorded"}


# ── Alerts / Proactive intelligence endpoints ─────────────────────────────────

@app.get("/alerts")
def list_alerts(org_id: str = "org-1", severity: Optional[str] = None) -> list[dict]:
    from dataclasses import asdict
    return [asdict(a) for a in service.scanner.list_alerts(org_id=org_id, severity=severity)]


@app.post("/alerts/{alert_id}/acknowledge")
def acknowledge_alert(alert_id: str) -> dict:
    service.scanner.acknowledge(alert_id)
    return {"status": "acknowledged"}


@app.post("/alerts/scan")
def run_proactive_scan(org_id: str = "org-1") -> dict:
    from dataclasses import asdict
    # Build scanner-compatible KPI dicts: scanner expects current_value, target_value, kpi_id, org_id
    kpi_status_list = service.kpis.kpi_status(org_id=org_id)
    scanner_kpis = [
        {
            "kpi_id": k.get("kpi_id", ""),
            "name": k.get("name", ""),
            "current_value": k.get("latest_value"),
            "target_value": k.get("target"),
            "unit": k.get("unit", ""),
            "org_id": k.get("org_id", org_id),
            "higher_is_better": True,
        }
        for k in kpi_status_list
    ]
    # Build scanner-compatible OKR dicts: scanner expects title, progress_pct, due_date, org_id
    objectives_raw = service.okrs.list_objectives(org_id=org_id)
    scanner_okrs = [
        {
            "title": o.name,
            "progress_pct": o.progress * 100,
            "due_date": None,
            "org_id": o.org_id,
            "objective_id": o.obj_id,
        }
        for o in objectives_raw
    ]
    budget_status = service.budgets.budget_status(org_id=org_id)

    kpi_alerts = service.scanner.scan_kpis(scanner_kpis)
    okr_alerts = service.scanner.scan_okrs(scanner_okrs)
    budget_alerts = service.scanner.scan_budget(budget_status)
    all_alerts = kpi_alerts + okr_alerts + budget_alerts

    return {
        "scanned": {"kpis": len(scanner_kpis), "objectives": len(scanner_okrs), "budget_categories": len(budget_status)},
        "alerts_generated": len(all_alerts),
        "alerts": [asdict(a) for a in all_alerts],
    }


# ── Financial modeling endpoints ──────────────────────────────────────────────

class ScenarioPayload(BaseModel):
    base_revenue: float
    base_costs: float
    optimistic_growth_pct: float = 0.30
    pessimistic_growth_pct: float = -0.15


class RunwayPayload(BaseModel):
    cash_on_hand: float
    monthly_burn: float
    current_mrr: float = 0
    mrr_growth_rate: float = 0.08


class UnitEconomicsPayload(BaseModel):
    arpu: float
    cac: float
    churn_rate: float
    gross_margin: float = 0.70


@app.post("/modeling/scenarios")
def run_scenarios(payload: ScenarioPayload) -> dict:
    from dataclasses import asdict
    scenarios = service.modeling.three_case_model(
        payload.base_revenue, payload.base_costs,
        payload.optimistic_growth_pct, payload.pessimistic_growth_pct)
    ev = service.modeling.expected_value(scenarios)
    return {"scenarios": [asdict(s) for s in scenarios], "expected_value": ev}


@app.post("/modeling/runway")
def calc_runway(payload: RunwayPayload) -> dict:
    from dataclasses import asdict
    return asdict(service.modeling.runway(
        payload.cash_on_hand, payload.monthly_burn,
        payload.current_mrr, payload.mrr_growth_rate))


@app.post("/modeling/unit-economics")
def unit_economics(payload: UnitEconomicsPayload) -> dict:
    return service.modeling.unit_economics(
        payload.arpu, payload.cac, payload.churn_rate, payload.gross_margin)


# ── Voice transcription endpoint ──────────────────────────────────────────────

@app.post("/voice/transcribe")
async def transcribe_voice(file: UploadFile = File(...)) -> dict:
    from dataclasses import asdict
    audio_bytes = await file.read()
    result = service.voice.transcribe_bytes(audio_bytes, filename=file.filename or "audio.mp3")
    return asdict(result)


# ── Weekly digest endpoint ────────────────────────────────────────────────────

@app.get("/digest/weekly")
def weekly_digest(org_id: str = "org-1") -> dict:
    from dataclasses import asdict
    # Normalize KPIs: digest expects current_value, target_value
    kpi_status_list = service.kpis.kpi_status(org_id=org_id)
    digest_kpis = [
        {
            "name": k.get("name", ""),
            "current_value": k.get("latest_value", 0) or 0,
            "target_value": k.get("target") or 1,
            "unit": k.get("unit", ""),
            "status": "on_track" if k.get("on_target") else "at_risk",
        }
        for k in kpi_status_list
    ]
    # Normalize OKRs: digest expects title, progress_pct (0-100), status
    digest_objectives = [
        {
            "title": o.name,
            "progress_pct": o.progress * 100,
            "status": o.status,
        }
        for o in service.okrs.list_objectives(org_id=org_id)
    ]
    alerts = [asdict(a) for a in service.scanner.list_alerts(org_id=org_id)]
    decisions = [asdict(d) for d in service.decisions.list_decisions(org_id=org_id, limit=10)]

    digest = service.digest.generate_weekly(digest_kpis, digest_objectives, alerts, decisions, org_id=org_id)
    result = asdict(digest)
    result["markdown"] = service.digest.digest_to_markdown(digest)
    return result


# ── QA Test Registry ──────────────────────────────────────────────────────────

class QATestCasePayload(BaseModel):
    title: str
    feature_area: str
    test_type: str
    org_id: str = "default"
    subfeature: str = ""
    description: str = ""
    preconditions: str = ""
    steps: list[str] = []
    expected_result: str = ""
    priority: str = "medium"
    severity_if_fails: str = "major"
    applies_to_agents: list[str] = []
    applies_to_ui_surfaces: list[str] = []
    release_blocker: bool = False
    status: str = "draft"
    created_by: str = "qa-specialist"
    linked_user_story_ids: list[str] = []
    linked_bug_ids: list[str] = []
    linked_workspace_ids: list[str] = []
    notes: str = ""
    tags: list[str] = []


class QASuitePayload(BaseModel):
    name: str
    suite_type: str
    org_id: str = "default"
    description: str = ""
    feature_areas: list[str] = []
    test_case_ids: list[str] = []
    generated_by_rule: str = ""
    owner: str = ""


class QASuiteGeneratePayload(BaseModel):
    name: str
    rule: str
    org_id: str = "default"
    owner: str = ""


class QARunPayload(BaseModel):
    suite_id: str
    title: str
    org_id: str = "default"
    environment: str = "development"
    triggered_by: str = "manual"
    run_type: str = "manual"
    notes: str = ""


class QAResultPayload(BaseModel):
    test_case_id: str
    result: str
    findings: str = ""
    reproduction_notes: str = ""
    severity: Optional[str] = None
    linked_bug_id: Optional[str] = None
    should_become_regression: bool = False
    tester: str = ""


class QABugPayload(BaseModel):
    title: str
    org_id: str = "default"
    severity: str = "major"
    category: str = "functional"
    area: str = ""
    repro_steps: list[str] = []
    expected_result: str = ""
    actual_result: str = ""
    impact: str = ""
    recommended_fix: str = ""
    release_blocker: bool = False
    linked_test_case_ids: list[str] = []
    linked_run_ids: list[str] = []
    owner: str = ""


class QAStoryPayload(BaseModel):
    title: str
    org_id: str = "default"
    user_story: str = ""
    context: str = ""
    acceptance_criteria: list[str] = []
    priority: str = "medium"
    dependencies: list[str] = []
    source_test_case_id: Optional[str] = None
    source_run_id: Optional[str] = None


# ─── Test Case Routes ────────────────────────────────────────────────────────

@app.get("/qa/tests")
def list_qa_tests(
    org_id: str = "default",
    feature_area: Optional[str] = None,
    test_type: Optional[str] = None,
    status: Optional[str] = None,
    priority: Optional[str] = None,
    release_blocker: Optional[bool] = None,
    search: Optional[str] = None,
    limit: int = 200,
    offset: int = 0,
) -> list[dict]:
    from dataclasses import asdict
    tests = service.qa.list_test_cases(
        org_id=org_id, feature_area=feature_area, test_type=test_type,
        status=status, priority=priority, release_blocker=release_blocker,
        search=search, limit=limit, offset=offset,
    )
    return [asdict(t) for t in tests]


@app.post("/qa/tests")
def create_qa_test(payload: QATestCasePayload) -> dict:
    from dataclasses import asdict
    tc = service.qa.create_test_case(**payload.model_dump())
    return asdict(tc)


@app.get("/qa/tests/{tc_id}")
def get_qa_test(tc_id: str) -> dict:
    from dataclasses import asdict
    tc = service.qa.get_test_case(tc_id)
    if not tc:
        raise HTTPException(status_code=404, detail="Test case not found")
    return asdict(tc)


@app.put("/qa/tests/{tc_id}")
def update_qa_test(tc_id: str, updates: dict[str, Any], updated_by: str = "qa-specialist") -> dict:
    from dataclasses import asdict
    tc = service.qa.update_test_case(tc_id, updates, updated_by=updated_by)
    if not tc:
        raise HTTPException(status_code=404, detail="Test case not found")
    return asdict(tc)


@app.post("/qa/tests/{tc_id}/deprecate")
def deprecate_qa_test(tc_id: str, updated_by: str = "qa-specialist") -> dict:
    ok = service.qa.deprecate_test_case(tc_id, updated_by=updated_by)
    if not ok:
        raise HTTPException(status_code=404, detail="Test case not found")
    return {"status": "deprecated", "tc_id": tc_id}


@app.post("/qa/tests/{tc_id}/clone")
def clone_qa_test(tc_id: str, new_title: Optional[str] = None, created_by: str = "qa-specialist") -> dict:
    from dataclasses import asdict
    tc = service.qa.clone_test_case(tc_id, new_title=new_title, created_by=created_by)
    if not tc:
        raise HTTPException(status_code=404, detail="Test case not found")
    return asdict(tc)


# ─── Suite Routes ────────────────────────────────────────────────────────────

@app.get("/qa/suites")
def list_qa_suites(org_id: str = "default", status: Optional[str] = None) -> list[dict]:
    from dataclasses import asdict
    return [asdict(s) for s in service.qa.list_suites(org_id=org_id, status=status)]


@app.post("/qa/suites")
def create_qa_suite(payload: QASuitePayload) -> dict:
    from dataclasses import asdict
    suite = service.qa.create_suite(**payload.model_dump())
    return asdict(suite)


@app.get("/qa/suites/{suite_id}")
def get_qa_suite(suite_id: str) -> dict:
    from dataclasses import asdict
    suite = service.qa.get_suite(suite_id)
    if not suite:
        raise HTTPException(status_code=404, detail="Suite not found")
    # Enrich with full test case objects
    result = asdict(suite)
    result["test_cases"] = [
        asdict(tc) for tc in [service.qa.get_test_case(tid) for tid in suite.test_case_ids]
        if tc is not None
    ]
    return result


@app.put("/qa/suites/{suite_id}")
def update_qa_suite(suite_id: str, updates: dict[str, Any]) -> dict:
    from dataclasses import asdict
    suite = service.qa.update_suite(suite_id, updates)
    if not suite:
        raise HTTPException(status_code=404, detail="Suite not found")
    return asdict(suite)


@app.post("/qa/suites/generate")
def generate_qa_suite(payload: QASuiteGeneratePayload) -> dict:
    from dataclasses import asdict
    suite = service.qa.generate_suite_from_rules(
        name=payload.name, rule=payload.rule,
        org_id=payload.org_id, owner=payload.owner,
    )
    return asdict(suite)


# ─── Run Routes ──────────────────────────────────────────────────────────────

@app.get("/qa/runs")
def list_qa_runs(org_id: str = "default", suite_id: Optional[str] = None, limit: int = 50) -> list[dict]:
    from dataclasses import asdict
    return [asdict(r) for r in service.qa.list_runs(org_id=org_id, suite_id=suite_id, limit=limit)]


@app.post("/qa/runs")
def create_qa_run(payload: QARunPayload) -> dict:
    from dataclasses import asdict
    run = service.qa.create_test_run(**payload.model_dump())
    return asdict(run)


@app.get("/qa/runs/{run_id}")
def get_qa_run(run_id: str) -> dict:
    from dataclasses import asdict
    run = service.qa.get_run(run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    result = asdict(run)
    result["results"] = [asdict(r) for r in service.qa.list_test_results(run_id)]
    return result


@app.post("/qa/runs/{run_id}/results")
def store_qa_result(run_id: str, payload: QAResultPayload) -> dict:
    from dataclasses import asdict
    result = service.qa.store_test_result(run_id=run_id, **payload.model_dump())
    return asdict(result)


@app.post("/qa/runs/{run_id}/complete")
def complete_qa_run(run_id: str, summary: str = "", notes: str = "") -> dict:
    from dataclasses import asdict
    run = service.qa.complete_test_run(run_id, summary=summary, notes=notes)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    return asdict(run)


# ─── Bug Report Routes ───────────────────────────────────────────────────────

@app.get("/qa/bugs")
def list_qa_bugs(org_id: str = "default", status: Optional[str] = None, area: Optional[str] = None) -> list[dict]:
    from dataclasses import asdict
    return [asdict(b) for b in service.qa.list_bug_reports(org_id=org_id, status=status, area=area)]


@app.post("/qa/bugs")
def create_qa_bug(payload: QABugPayload) -> dict:
    from dataclasses import asdict
    bug = service.qa.create_bug_report(**payload.model_dump())
    return asdict(bug)


@app.put("/qa/bugs/{bug_id}")
def update_qa_bug(bug_id: str, updates: dict[str, Any]) -> dict:
    from dataclasses import asdict
    bug = service.qa.update_bug_report(bug_id, updates)
    if not bug:
        raise HTTPException(status_code=404, detail="Bug report not found")
    return asdict(bug)


# ─── User Story Routes ───────────────────────────────────────────────────────

@app.get("/qa/stories")
def list_qa_stories(org_id: str = "default", status: Optional[str] = None) -> list[dict]:
    from dataclasses import asdict
    return [asdict(s) for s in service.qa.list_user_story_candidates(org_id=org_id, status=status)]


@app.post("/qa/stories")
def create_qa_story(payload: QAStoryPayload) -> dict:
    from dataclasses import asdict
    story = service.qa.create_user_story_candidate(**payload.model_dump())
    return asdict(story)


# ─── Coverage Gap Analysis ───────────────────────────────────────────────────

@app.get("/qa/coverage")
def get_qa_coverage(org_id: str = "default") -> dict:
    from dataclasses import asdict
    report = service.qa.analyze_coverage_gaps(org_id=org_id)
    result = asdict(report)
    result["gaps"] = [asdict(g) for g in report.gaps]
    return result


# ─── Registry Summary ────────────────────────────────────────────────────────

@app.get("/qa/summary")
def get_qa_summary(org_id: str = "default") -> dict:
    return service.qa.get_registry_summary(org_id=org_id)


# ─── Seed Templates ──────────────────────────────────────────────────────────

@app.post("/qa/templates/seed")
def seed_qa_templates(org_id: str = "default") -> dict:
    """Force re-seed default templates (useful after schema changes)."""
    service.qa._seed_templates()
    summary = service.qa.get_registry_summary(org_id=org_id)
    return {"status": "seeded", "total": summary["total"]}


# ─── Regression Candidates ───────────────────────────────────────────────────

@app.get("/qa/regression-candidates")
def get_regression_candidates(org_id: str = "default") -> list[dict]:
    from dataclasses import asdict
    return [asdict(r) for r in service.qa.get_regression_candidates(org_id=org_id)]


# ─────────────────────────────────────────────────────────────────────────────

@app.post("/admin/runtime/reload")
def admin_runtime_reload(request: Request) -> dict:
    admin_auth.require(request)
    global service
    service = FridayService()
    return {"status": "reloaded"}
